from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.17)
    section.right_margin  = Cm(3.17)

# ── Helper: set paragraph font ────────────────────────────────────────────────
def fmt(para, bold=False, size=12, color=None, italic=False):
    for run in para.runs:
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)

# ── Helper: add a styled heading ──────────────────────────────────────────────
def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)   # dark navy
    run.font.size = Pt(15 if level == 1 else 13)
    run.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

# ── Helper: normal paragraph ──────────────────────────────────────────────────
def add_para(text, bold=False, size=11.5, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    return p

# ── Helper: bullet point ──────────────────────────────────────────────────────
def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p

# ── Helper: simple table ──────────────────────────────────────────────────────
def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Dark navy background
        tc   = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1F3A6E')
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = val
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)
        # Alternate row shading
        if r_idx % 2 == 0:
            for cell in cells:
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  'E8ECF5')
                tcPr.append(shd)

    doc.add_paragraph()  # spacing after table
    return table

# ── Helper: horizontal rule ───────────────────────────────────────────────────
def add_hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3A6E')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(40)
r = title.add_run('PROJECT REPORT')
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Ride-Sharing Demand Forecasting')
r2.bold = True
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

doc.add_paragraph()
doc.add_paragraph()

meta_lines = [
    'Submitted as part of Internship Training',
    '',
    'Technologies Used: Python | Scikit-learn | Pandas | NumPy | Matplotlib | Seaborn',
    '',
    'Development Environment: Google Colab | VS Code',
]
for line in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  VI. PROJECT WORK
# ══════════════════════════════════════════════════════════════════════════════
add_heading('VI. PROJECT WORK: RIDE-SHARING DEMAND FORECASTING', level=1)
add_hr()

# 1. Introduction
add_heading('1. Introduction to the Project', level=2)
add_para(
    'The project focuses on predicting ride-sharing demand using machine learning techniques. '
    'It aims to analyze historical data and forecast future demand.'
)
add_para(
    'This project was developed as part of the internship to apply the concepts learned throughout '
    'the training in a real-world scenario. Ride-sharing demand forecasting is an important problem '
    'in the transportation domain, as it directly impacts service efficiency and customer experience. '
    'By using historical ride data, the project attempts to identify patterns and trends that can be '
    'used to predict future demand. The implementation involved working with real datasets, performing '
    'analysis, and building machine learning models to generate accurate predictions.'
)
add_hr()

# 2. Problem Statement
add_heading('2. Problem Statement', level=2)
add_para('Ride-sharing platforms face challenges in managing supply and demand. Accurate demand prediction helps in:')
add_bullet('Reducing waiting time')
add_bullet('Optimizing driver allocation')
add_bullet('Improving customer satisfaction')
add_para(
    'In real-world scenarios, demand for rides fluctuates based on time, location, and external factors. '
    'Without proper prediction, there can be situations where there are either too many drivers with fewer '
    'customers or too many customers with fewer drivers. This imbalance affects both service providers and '
    'users. Therefore, building a reliable demand prediction system is essential for improving operational '
    'efficiency and ensuring better service delivery.'
)
add_hr()

# 3. Objectives
add_heading('3. Objectives', level=2)
add_bullet('Build a predictive model for demand forecasting')
add_bullet('Analyze trends in ride usage')
add_bullet('Improve accuracy using machine learning')
add_para(
    'The main objective of the project was to develop a model that can accurately predict ride demand '
    'based on historical data. Additionally, the project aimed to understand usage patterns such as peak '
    'hours and high-demand locations. Another key objective was to compare different machine learning '
    'models and select the one that provides the best performance in terms of accuracy and reliability.'
)
add_hr()

# 4. Dataset Description
add_heading('4. Dataset Description', level=2)
add_table(
    ['Feature', 'Description'],
    [
        ['Date and Time',       'Timestamp of each ride request'],
        ['Location Details',    'Pickup and drop-off area information'],
        ['Ride Counts',         'Target variable for prediction'],
        ['Additional Factors',  'Other influencing variables'],
    ]
)
add_para(
    'The dataset used in this project consisted of ride-sharing data containing information about when '
    'and where rides were requested. The date and time feature helped in extracting time-based patterns, '
    'while location details provided insights into demand across different areas. Ride counts served as '
    'the target variable for prediction. Additional features, if present, contributed to improving the '
    'overall accuracy of the model. Understanding the dataset structure was an important step before '
    'proceeding with preprocessing and analysis.'
)
add_hr()

# 5. Data Preprocessing
add_heading('5. Data Preprocessing', level=2)
add_para('Data preprocessing was a crucial step in preparing the dataset for modeling. The following steps were performed:')
steps = [
    ('Handling Missing Values – ', 'Missing values were handled using appropriate techniques to ensure data completeness.'),
    ('Feature Engineering – ',     'New meaningful variables were created from existing data.'),
    ('Encoding Categorical Variables – ', 'Categorical variables were converted into numerical form using encoding techniques so that they could be used by machine learning algorithms.'),
    ('Feature Scaling – ',         'Numerical features were scaled to maintain uniformity and improve model performance.'),
]
for i, (bold_part, normal_part) in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(bold_part)
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(normal_part)
    r2.font.size = Pt(11)
add_para('These steps ensured that the dataset was clean, consistent, and suitable for analysis.')
add_hr()

# 6. EDA
add_heading('6. Exploratory Data Analysis', level=2)
add_para('EDA revealed key patterns and insights from the dataset:')
add_bullet('Peak demand hours',       'Peak Demand Hours – ')
add_bullet('Daily and weekly trends', 'Daily and Weekly Trends – ')
add_bullet('Location-based variations','Location-Based Variations – ')
add_para(
    'Exploratory Data Analysis helped in understanding the dataset more deeply. By visualizing the data, '
    'it was possible to identify peak hours during which ride demand was highest. Daily and weekly trends '
    'showed how demand changes over time, such as higher demand on weekends or during specific hours of '
    'the day. Location-based analysis revealed areas with consistently high or low demand. These insights '
    'were important for selecting relevant features and improving model performance.'
)
add_hr()

# 7. Feature Engineering
add_heading('7. Feature Engineering', level=2)
add_table(
    ['Feature', 'Purpose'],
    [
        ['Hour of the Day',      'Captures time-based demand variations'],
        ['Day of the Week',      'Highlights weekday vs. weekend differences'],
        ['Peak Hour Indicators', 'Marks high-demand time periods'],
    ]
)
add_para(
    'Feature engineering enhanced the dataset by creating additional variables that capture important '
    'patterns. Extracting the hour of the day helped in identifying time-based demand variations, while '
    'the day of the week highlighted differences between weekdays and weekends. Peak hour indicators were '
    'used to mark high-demand periods, which improved the model\'s ability to predict demand more '
    'accurately. These engineered features played a significant role in boosting model performance.'
)
add_hr()

# 8. Model Building
add_heading('8. Model Building', level=2)
add_table(
    ['Model', 'Description'],
    [
        ['Linear Regression', 'Baseline model for understanding basic relationships in the data'],
        ['Decision Tree',     'Captures non-linear patterns and provides interpretability'],
        ['Random Forest',     'Ensemble method combining multiple decision trees to reduce overfitting'],
    ]
)
add_para(
    'Multiple machine learning models were implemented to compare their performance. Linear Regression '
    'was used as a baseline model to understand basic relationships in the data. Decision Tree was applied '
    'to capture non-linear patterns and provide interpretability. Random Forest, being an ensemble method, '
    'combined multiple decision trees to improve accuracy and reduce overfitting. Each model was trained '
    'using the prepared dataset and evaluated to determine its effectiveness.'
)
add_hr()

# 9. Model Evaluation
add_heading('9. Model Evaluation', level=2)
add_table(
    ['Metric', 'Value'],
    [
        ['Mean Absolute Error (MAE)', '37.83'],
        ['R² Score',                  '0.89'],
    ]
)
add_para(
    'Model evaluation was carried out using appropriate performance metrics. Mean Absolute Error (MAE) '
    'measured the average difference between predicted and actual values, indicating prediction accuracy. '
    'The R² score measured how well the model explains the variance in the data. An R² score of 0.89 '
    'indicates that the model performs well and captures most of the patterns in the dataset. These '
    'metrics helped in comparing models and selecting the best-performing one.'
)
add_hr()

# 10. Result Analysis
add_heading('10. Result Analysis', level=2)
p = doc.add_paragraph()
r = p.add_run(
    'The Random Forest model performed better compared to other models due to its ability to handle '
    'non-linear relationships.'
)
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
add_para(
    'The results showed that Random Forest outperformed Linear Regression and Decision Tree in terms of '
    'accuracy and reliability. This is because Random Forest reduces overfitting by combining multiple '
    'trees and captures complex relationships within the data effectively. The model was able to predict '
    'demand with high accuracy, making it suitable for real-world applications. Overall, the project '
    'successfully demonstrated how machine learning techniques can be used to solve practical problems '
    'in the transportation domain.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  VII. TOOLS AND TECHNIQUES
# ══════════════════════════════════════════════════════════════════════════════
add_heading('VII. TOOLS AND TECHNIQUES USED', level=1)
add_hr()

add_table(
    ['Tool / Library', 'Purpose'],
    [
        ['Python',        'Primary programming language'],
        ['Pandas',        'Data manipulation and analysis'],
        ['NumPy',         'Numerical computations and array handling'],
        ['Matplotlib',    'Basic data visualization and plotting'],
        ['Seaborn',       'Advanced and visually appealing graphs'],
        ['Scikit-learn',  'Machine learning model implementation and evaluation'],
        ['Google Colab',  'Cloud-based development environment'],
        ['VS Code',       'Code editor for writing and managing scripts'],
    ]
)

tools = [
    ('1. Python',
     'Python was the primary programming language used throughout the project due to its simplicity, '
     'readability, and wide range of libraries available for data science. It enabled efficient '
     'implementation of data analysis, preprocessing, and machine learning tasks.'),
    ('2. Pandas, NumPy',
     'Pandas was used for data manipulation and analysis, allowing easy handling of datasets through '
     'operations like filtering, grouping, and transformation. NumPy was used for numerical computations '
     'and handling arrays, providing support for mathematical operations required during preprocessing '
     'and modeling.'),
    ('3. Matplotlib, Seaborn',
     'These libraries were used for data visualization. Matplotlib helped in creating basic plots, while '
     'Seaborn provided advanced and more visually appealing graphs. They were essential for performing '
     'exploratory data analysis and understanding patterns in the dataset.'),
    ('4. Scikit-learn',
     'Scikit-learn was used for implementing machine learning models such as Linear Regression, Decision '
     'Tree, and Random Forest. It also provided tools for model evaluation, data splitting, and '
     'preprocessing, making it a key library for the project.'),
    ('5. Google Colab',
     'Google Colab was used as the primary development environment for coding and running the project. '
     'Being cloud-based, it provided sufficient computational power and eliminated the need for '
     'high-end local systems.'),
    ('6. VS Code',
     'Visual Studio Code was used as a code editor for writing and managing scripts. It provided a '
     'professional development environment with features like debugging and extensions that improved '
     'coding efficiency.'),
]
for heading, body in tools:
    add_heading(heading, level=2)
    add_para(body)
add_hr()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  VIII. OBSERVATIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('VIII. OBSERVATIONS', level=1)
add_hr()

add_table(
    ['Observation', 'Details'],
    [
        ['Data quality directly affects model accuracy',
         'Clean, well-processed data resulted in better predictions; noisy or incomplete data reduced model performance.'],
        ['Feature engineering improves performance',
         'Creating meaningful features from existing data helped the model understand patterns more effectively.'],
        ['Visualization helps in better understanding',
         'Graphical representations provided clear insights for better decision-making during analysis and model building.'],
    ]
)

obs = [
    ('Data Quality – ',
     'The quality of data plays a major role in determining the accuracy of machine learning models. '
     'Clean, well-processed data resulted in better predictions, whereas noisy or incomplete data '
     'reduced model performance significantly.'),
    ('Feature Engineering – ',
     'Creating meaningful features from existing data helped the model understand patterns more '
     'effectively. Proper feature engineering enhanced the predictive capability of the model without '
     'changing the algorithm.'),
    ('Visualization – ',
     'Visualization techniques made it easier to interpret data and identify important trends. '
     'Graphical representations provided clear insights, which helped in making better decisions '
     'during analysis and model building.'),
]
for bold_part, normal_part in obs:
    add_bullet(normal_part, bold_prefix=bold_part)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r'c:\Users\hp\Downloads\heart+disease\Ride_Sharing_Demand_Forecasting_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
