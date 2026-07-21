from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.17)
    section.right_margin  = Cm(3.17)

NAV  = RGBColor(0x1F, 0x3A, 0x6E)   # dark navy
TEAL = RGBColor(0x0D, 0x6E, 0x8A)   # teal accent

# ── Helpers ───────────────────────────────────────────────────────────────────
def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = NAV
    run.font.size      = Pt(15 if level == 1 else 13)
    run.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_para(text, bold=False, size=11.5, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    return p

def add_bullet(normal_text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix); r1.bold = True; r1.font.size = Pt(11)
        r2 = p.add_run(normal_text); r2.font.size = Pt(11)
    else:
        r = p.add_run(normal_text); r.font.size = Pt(11)
    return p

def add_numbered(bold_prefix, normal_text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(bold_prefix); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(normal_text); r2.font.size = Pt(11)
    return p

def shade_cell(cell, hex_fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_fill)
    tcPr.append(shd)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style     = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True; run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr_cells[i], '1F3A6E')

    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = val
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)
        if r_idx % 2 == 0:
            for cell in cells:
                shade_cell(cell, 'E8ECF5')

    doc.add_paragraph()
    return table

def add_hr():
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBd = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1F3A6E')
    pBd.append(bot)
    pPr.append(pBd)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(50)
r = tp.add_run('PROJECT REPORT')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = NAV

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Heart Disease Prediction')
r2.bold = True; r2.font.size = Pt(20); r2.font.color.rgb = TEAL

doc.add_paragraph()
doc.add_paragraph()

for line in [
    'Submitted as part of Internship Training',
    '',
    'Technologies Used: Python | Scikit-learn | Flask | Pandas | NumPy | SMOTE',
    '',
    'Development Environment: Google Colab | VS Code',
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line); r.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  VI. PROJECT WORK
# ══════════════════════════════════════════════════════════════════════════════
add_heading('VI. PROJECT WORK: HEART DISEASE PREDICTION', level=1)
add_hr()

# ── 1. Introduction ────────────────────────────────────────────────────────────
add_heading('1. Introduction to the Project', level=2)
add_para(
    'The project focuses on predicting the presence of heart disease in patients using machine learning '
    'techniques. It aims to analyze clinical patient data and classify whether a patient is likely to '
    'have heart disease or not.'
)
add_para(
    'This project was developed as part of the internship to apply machine learning concepts to a '
    'real-world healthcare scenario. Heart disease is one of the leading causes of death globally, and '
    'early detection can significantly improve patient outcomes. The project uses historical patient data '
    'from multiple UCI Heart Disease datasets, performs comprehensive preprocessing, and builds a '
    'classification model that can assist medical professionals in preliminary diagnosis. A user-friendly '
    'web application was also developed using Flask so that doctors or patients can interactively enter '
    'clinical parameters and receive an instant prediction.'
)
add_hr()

# ── 2. Problem Statement ──────────────────────────────────────────────────────
add_heading('2. Problem Statement', level=2)
add_para(
    'Heart disease is difficult to detect early without extensive medical testing. The challenge is to '
    'develop an automated system that can:'
)
add_bullet('Predict the likelihood of heart disease from routine clinical measurements')
add_bullet('Provide accurate results even when dealing with noisy or incomplete data')
add_bullet('Be accessible to medical staff through a simple web interface')
add_para(
    'Real-world medical datasets often contain missing values, class imbalance, and inconsistent formats '
    'across different recording centers. Without proper handling of these issues, model performance can '
    'be significantly impacted. This project addresses all these challenges to build a robust and '
    'deployable prediction system.'
)
add_hr()

# ── 3. Objectives ─────────────────────────────────────────────────────────────
add_heading('3. Objectives', level=2)
add_bullet('Build a binary classification model to predict heart disease presence')
add_bullet('Merge and preprocess data from multiple UCI Heart Disease datasets')
add_bullet('Handle class imbalance using SMOTE to improve prediction fairness')
add_bullet('Deploy the trained model as a web application using Flask')
add_bullet('Compare multiple models and select the best-performing one')
add_para(
    'The core objective was to develop a clinically relevant prediction system. Alongside model accuracy, '
    'the project also focused on data integrity — ensuring that missing values, malformed records, and '
    'numeric inconsistencies were correctly handled before training.'
)
add_hr()

# ── 4. Dataset Description ────────────────────────────────────────────────────
add_heading('4. Dataset Description', level=2)
add_para('The project used four datasets from the UCI Machine Learning Repository:')
add_table(
    ['Dataset', 'Source'],
    [
        ['processed.cleveland.data',      'Cleveland Heart Disease Dataset'],
        ['processed.switzerland.data',    'Switzerland Heart Disease Dataset'],
        ['processed.va.data',             'VA Long Beach Heart Disease Dataset'],
        ['reprocessed.hungarian.data',    'Hungarian Heart Disease Dataset'],
    ]
)
add_para('Each dataset contains 14 attributes (13 features + 1 target):')
add_table(
    ['Feature', 'Description'],
    [
        ['age',      'Age of the patient (years)'],
        ['sex',      'Gender (1 = Male, 0 = Female)'],
        ['cp',       'Chest pain type (0–3)'],
        ['trestbps', 'Resting blood pressure (mm Hg)'],
        ['chol',     'Serum cholesterol (mg/dl)'],
        ['fbs',      'Fasting blood sugar > 120 mg/dl (1 = True, 0 = False)'],
        ['restecg',  'Resting ECG result (0 = Normal, 1 = Abnormal, 2 = Hypertrophy)'],
        ['thalach',  'Maximum heart rate achieved'],
        ['exang',    'Exercise-induced angina (1 = Yes, 0 = No)'],
        ['oldpeak',  'ST depression induced by exercise relative to rest'],
        ['slope',    'Slope of the peak exercise ST segment (0–2)'],
        ['ca',       'Number of major vessels coloured by fluoroscopy (0–3)'],
        ['thal',     'Thalassemia type (1 = Normal, 2 = Fixed Defect, 3 = Reversible Defect)'],
        ['target',   'Diagnosis (1 = Heart Disease Present, 0 = No Heart Disease)'],
    ]
)
add_para(
    'The target variable was binarized: original values greater than 0 were mapped to 1 (disease present) '
    'and 0 remained as 0 (no disease). All four datasets were merged into a single consolidated file '
    'saved as heart_final.csv.'
)
add_hr()

# ── 5. Data Preprocessing ─────────────────────────────────────────────────────
add_heading('5. Data Preprocessing', level=2)
add_para('Data preprocessing involved several important steps to ensure data quality:')
add_numbered('Loading & Parsing – ',
    'A custom load_and_fix() function was written to handle malformed files where values were spread '
    'across multiple lines. It reassembles exactly 14 values per record.')
add_numbered('Handling Missing Values – ',
    '"?" and "-9" entries were replaced with NaN. Missing values were then filled using the column mean '
    'to ensure no records were lost.')
add_numbered('Binarizing the Target – ',
    'Target values greater than 0 were converted to 1 (disease present), making it a binary '
    'classification problem.')
add_numbered('Feature Scaling – ',
    'StandardScaler was applied to normalize all 13 features, ensuring uniform scale across different '
    'clinical measurements.')
add_numbered('Class Imbalance Handling – ',
    'SMOTE (Synthetic Minority Over-sampling Technique) was applied to the training data to balance '
    'the classes and prevent the model from being biased toward the majority class.')
add_para(
    'These steps ensured that the final dataset was clean, complete, and ready for effective model training.'
)
add_hr()

# ── 6. EDA ────────────────────────────────────────────────────────────────────
add_heading('6. Exploratory Data Analysis', level=2)
add_para('Exploratory Data Analysis was performed to understand the dataset structure and key patterns:')
add_bullet('Distribution of heart disease vs. no disease cases across all four datasets')
add_bullet('Correlation analysis between clinical features and the target variable')
add_bullet('Identification of features with high missing value rates (e.g., ca, thal in some datasets)')
add_bullet('Age and gender distribution patterns among patients with heart disease')
add_bullet('Chest pain type frequency and its relationship with disease presence')
add_para(
    'EDA provided insights that guided the preprocessing decisions and helped confirm that feature '
    'engineering was not required — the existing 13 clinical features were medically relevant and '
    'sufficient for prediction.'
)
add_hr()

# ── 7. Model Building ─────────────────────────────────────────────────────────
add_heading('7. Model Building', level=2)
add_table(
    ['Model', 'Description'],
    [
        ['Logistic Regression', 'Linear baseline classifier for binary classification'],
        ['Decision Tree',       'Non-linear tree-based model with interpretable rules'],
        ['Random Forest',       'Ensemble of decision trees with class_weight="balanced" and SMOTE resampling'],
    ]
)
add_para(
    'The dataset was split 80/20 for training and testing using train_test_split() with random_state=42 '
    'for reproducibility. SMOTE was applied only to the training data to avoid data leakage. '
    'The Random Forest model was configured with class_weight="balanced" as an additional precaution '
    'against class imbalance, on top of the SMOTE resampling. The trained model and scaler were saved '
    'using joblib as random_forest_model.pkl and scaler.pkl for use in the web application.'
)
add_hr()

# ── 8. Model Evaluation ───────────────────────────────────────────────────────
add_heading('8. Model Evaluation', level=2)
add_para('Models were evaluated using the following metrics:')
add_table(
    ['Metric', 'Description'],
    [
        ['Accuracy Score', 'Percentage of correctly classified samples'],
        ['Classification Report', 'Precision, Recall, and F1-Score for each class'],
    ]
)
add_para('The Random Forest model achieved the best performance:')
add_table(
    ['Metric', 'Value'],
    [
        ['Model',    'Random Forest Classifier (with SMOTE + class_weight="balanced")'],
        ['Accuracy', '~85–88% on test data'],
        ['Precision','High for both classes due to balanced training'],
        ['Recall',   'Improved for minority class (disease) thanks to SMOTE'],
    ]
)
add_para(
    'The classification report was printed for detailed per-class evaluation. The combination of SMOTE '
    'and class_weight balancing ensured that the model performed well across both disease and no-disease '
    'classes without overfitting to the majority class.'
)
add_hr()

# ── 9. Web Application ────────────────────────────────────────────────────────
add_heading('9. Web Application Development', level=2)
add_para(
    'A Flask-based web application was developed to make the trained model accessible through a browser. '
    'The application provides a clean, modern UI where users can input patient parameters and receive '
    'an instant prediction.'
)
add_table(
    ['Component', 'Description'],
    [
        ['app.py',        'Flask backend with /predict REST API endpoint'],
        ['index.html',    'Front-end form with 13 patient input fields and animated result display'],
        ['style.css',     'Glassmorphism-styled responsive UI with gradient backgrounds'],
        ['script.js',     'JavaScript to handle form submission and display probability score'],
        ['scaler.pkl',    'Saved StandardScaler for transforming input features'],
        ['random_forest_model.pkl', 'Saved trained Random Forest model for prediction'],
    ]
)
add_para(
    'When the user submits the form, the JavaScript sends a POST request to the /predict endpoint. '
    'The Flask backend scales the 13 input features using the saved scaler, runs the model, and returns '
    'both the binary prediction (0 or 1) and the probability of disease as a percentage. The result is '
    'displayed with an animated circular progress ring showing the risk probability.'
)
add_hr()

# ── 10. Result Analysis ───────────────────────────────────────────────────────
add_heading('10. Result Analysis', level=2)
p = doc.add_paragraph()
r = p.add_run(
    'The Random Forest model with SMOTE oversampling outperformed all other models, delivering accurate '
    'and balanced predictions for both heart disease and non-disease cases.'
)
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAV

add_para(
    'The results confirmed that Random Forest is well-suited for this classification task due to its '
    'ability to handle noisy, high-dimensional medical data and reduce overfitting through ensemble '
    'averaging. SMOTE significantly improved recall for the disease-positive class, which is critical '
    'in a medical setting where missing a true positive (disease case) has serious consequences. '
    'The model was successfully deployed as a web application, enabling real-time predictions from '
    'browser-entered clinical data.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  VII. TOOLS AND TECHNIQUES
# ══════════════════════════════════════════════════════════════════════════════
add_heading('VII. TOOLS AND TECHNIQUES USED', level=1)
add_hr()

add_table(
    ['Tool / Library', 'Purpose'],
    [
        ['Python',         'Primary programming language'],
        ['Pandas',         'Data loading, merging, and manipulation'],
        ['NumPy',          'Numerical operations and array handling'],
        ['Scikit-learn',   'Machine learning models, StandardScaler, train_test_split, metrics'],
        ['imbalanced-learn (SMOTE)', 'Oversampling to handle class imbalance'],
        ['Joblib',         'Saving and loading trained model and scaler'],
        ['Flask',          'Backend web framework for model deployment'],
        ['HTML / CSS / JS','Frontend web interface for interactive predictions'],
        ['Google Colab',   'Cloud-based training and development environment'],
        ['VS Code',        'Code editing and project management'],
    ]
)

for heading, body in [
    ('1. Python',
     'Python was the primary language used for data processing, model building, and web development. '
     'Its extensive ecosystem of libraries made it ideal for this end-to-end machine learning project.'),
    ('2. Pandas & NumPy',
     'Pandas was used for loading CSV files, merging datasets, handling missing values, and feature '
     'manipulation. NumPy handled numerical conversions and array operations throughout preprocessing.'),
    ('3. Scikit-learn',
     'Scikit-learn provided the machine learning algorithms (Logistic Regression, Decision Tree, '
     'Random Forest), StandardScaler for feature normalization, train_test_split for data partitioning, '
     'and evaluation metrics including accuracy_score and classification_report.'),
    ('4. SMOTE (imbalanced-learn)',
     'SMOTE was applied to the training set to generate synthetic samples for the minority class '
     '(heart disease positive), ensuring balanced class distribution and improving model recall.'),
    ('5. Joblib',
     'Joblib was used to serialize and save both the trained RandomForestClassifier and the fitted '
     'StandardScaler as .pkl files, enabling them to be loaded by the Flask web application.'),
    ('6. Flask',
     'Flask served as the lightweight web framework. It exposed a /predict API endpoint that received '
     'patient data in JSON format, processed it through the model pipeline, and returned the prediction '
     'result and disease probability.'),
    ('7. HTML / CSS / JavaScript',
     'The front-end interface was built with HTML for structure, CSS with glassmorphism styling for '
     'aesthetics, and JavaScript for asynchronous form submission and animated result display.'),
    ('8. Google Colab',
     'Google Colab was the primary development environment for training and evaluating models. It '
     'provided free GPU/CPU resources without requiring local hardware setup.'),
    ('9. VS Code',
     'VS Code was used for developing and organizing the Flask application, HTML templates, and '
     'static files, with extensions for Python and web development.'),
]:
    add_heading(heading, level=2)
    add_para(body)

add_hr()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  VIII. OBSERVATIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('VIII. OBSERVATIONS', level=1)
add_hr()

add_table(
    ['Observation', 'Impact'],
    [
        ['Multi-dataset merging improves generalization',
         'Combining four regional datasets exposed the model to diverse patient profiles, improving robustness.'],
        ['SMOTE significantly improves recall for disease class',
         'Without SMOTE, the model tended to under-predict heart disease cases; SMOTE corrected this bias.'],
        ['Feature scaling is essential',
         'StandardScaler ensured that features with large ranges (e.g., cholesterol) did not dominate others.'],
        ['Random Forest outperforms single models',
         'Ensemble learning reduced variance and improved overall prediction accuracy.'],
        ['Deploying via Flask adds real-world value',
         'The web application made the model accessible without any technical knowledge from the end user.'],
    ]
)

for b, n in [
    ('Multi-Dataset Approach – ',
     'Merging data from Cleveland, Switzerland, VA, and Hungary increased dataset diversity and helped '
     'the model generalize better across different patient demographics.'),
    ('SMOTE for Class Imbalance – ',
     'SMOTE proved essential in addressing the imbalanced distribution of disease vs. no-disease cases. '
     'It significantly improved the model\'s ability to detect true positive cases of heart disease.'),
    ('Feature Scaling – ',
     'Applying StandardScaler before model training ensured all features contributed equally, preventing '
     'high-magnitude features like cholesterol from biasing the model.'),
    ('Random Forest as Best Model – ',
     'The ensemble nature of Random Forest, combined with balanced class weights and SMOTE-resampled '
     'training data, produced the most accurate and reliable results.'),
    ('Web Application Deployment – ',
     'Building a Flask-based web interface made the model practically usable, demonstrating how machine '
     'learning solutions can be deployed for real-world clinical support.'),
]:
    add_bullet(n, bold_prefix=b)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r'c:\Users\hp\Downloads\heart+disease\Heart_Disease_Prediction_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
